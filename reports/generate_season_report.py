#!/usr/bin/env python3
"""生成 T 雷达站 RMUC2026 东部赛区赛季总结报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_table(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for para in header_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(header_cells[i], header_color)

    # 数据行
    for r, row_data in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, cell_text in enumerate(row_data):
            row_cells[c].text = str(cell_text)
            for para in row_cells[c].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8.5)
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
            # 交替行底色
            if r % 2 == 1:
                set_cell_shading(row_cells[c], 'E8F0FE')

    # 设置列宽
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    return table

def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_para(doc, text, bold=False, size=10.5, indent=False):
    """添加段落"""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = bold
    return para

def add_rich_para(doc, segments):
    """添加富文本段落，segments = [(text, bold, color), ...]"""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Cm(0.74)
    for text, bold, color in segments:
        run = para.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.bold = bold
        if color:
            run.font.color.rgb = color
    return para


def main():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('T 雷达站')
    run.font.size = Pt(36)
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('RMUC2026 东部赛区 · 赛季深度复盘')
    run.font.size = Pt(18)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('比赛日期：2026.05.20 — 2026.05.24\n报告生成：2026.06.27')
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ========== 目录页 ==========
    add_heading_styled(doc, '目录', level=1)
    toc_items = [
        '3.1.1  战术定位与赛场实战复盘',
        '3.1.2  核心需求与赛场数据闭环',
        '3.1.3  核心方案成败总结与经验沉淀',
        '3.1.4  人力资源与能力缺口复盘',
        '3.1.5  赛场故障与测试盲区复盘',
        '3.1.6  遗留问题与技术传承库',
    ]
    for item in toc_items:
        add_para(doc, item, size=12)
    doc.add_page_break()

    # ==========================================
    # 3.1.1 战术定位与赛场实战复盘
    # ==========================================
    add_heading_styled(doc, '3.1.1  战术定位与赛场实战复盘', level=1)

    add_para(doc, 'T 雷达站的核心战术定位是"全场多目标视觉-LiDAR 融合探测 + 裁判系统串口实时上报"。通过 6MP 海康相机（MV-CS060-10UC-PRO）+ Livox Mid-70 激光雷达，实现比赛场地内全部敌方机器人的实时定位与跟踪，将坐标通过裁判系统串口协议（cmd_id=0x0305，≤5Hz，48B/帧）上报至官方小地图，为己方操作手提供全场态势感知。')

    add_heading_styled(doc, '赛季实战总览', level=2)

    add_styled_table(doc,
        ['指标', '数值'],
        [
            ['比赛日期', '2026.05.20 — 05.24（5天）'],
            ['自录比赛包', '59 场'],
            ['有效对手专项包', '3 个（浙大1、浙大2、中国石油大学）'],
            ['总录包数据量', '~63 GB'],
            ['对手数据总量', '~7.5 GB'],
            ['最大单场数据', '13 GB（match_20260521_211546）'],
            ['代码规模', '~22,800 行（46 C++ + 18 Python）'],
            ['赛季构建次数', '84 次'],
        ],
        col_widths=[5, 11]
    )

    add_heading_styled(doc, '比赛日节奏与强度', level=2)

    add_styled_table(doc,
        ['日期', '场次', '数据量', '特点分析'],
        [
            ['5.20', '12 场', '~8.3 GB', '首日磨合，单场偏小（多数<1GB），凌晨仍有密集构建'],
            ['5.21', '28 场', '~29.5 GB', '★ 高强度对战日，14:00-21:00 连续作战，2场10GB+大包'],
            ['5.22', '6 场', '~8.2 GB', '下午-晚间集中对战，单场大（1.3-3.9 GB）'],
            ['5.23', '5 场', '~8.0 GB', '下午核心战，赛后集中修改融合算法'],
            ['5.24', '8 场', '~5.7 GB', '收官日，上午密集小场 + 最后两场2.6-2.7 GB'],
        ],
        col_widths=[2, 2, 3, 9]
    )

    add_heading_styled(doc, '异常包分析（空包/极小包）', level=2)

    add_para(doc, '在 59 场录包中，存在 8 场异常小包（≤68KB），表明启动后立即崩溃或被手动终止：')

    add_styled_table(doc,
        ['包名', '大小', '时间段', '推测原因'],
        [
            ['match_20260520_022741', '24 KB', '5.20 凌晨', '启动后即崩溃/手动终止'],
            ['match_20260521_093105', '4 KB', '5.21 早上', '测试启动-立即停止'],
            ['match_20260521_161152~163500', '16-24 KB ×6', '5.21 16:00-16:35', '★ 连续6次启动失败，35分钟内反复尝试'],
            ['match_20260521_201349', '68 KB', '5.21 晚间', '短暂运行后停止'],
        ],
        col_widths=[4.5, 2, 3.5, 6]
    )

    add_rich_para(doc, [
        ('★ 关键发现：', True, RGBColor(0xC0, 0x39, 0x2B)),
        ('5月21日 16:00-16:35 连续 6 场空包，说明该时段出现了系统性启动故障（可能是串口/LiDAR/相机之一未就绪），且当时未建立有效的"启动自检→阻止空包录制"机制。', False, None),
    ])

    # ==========================================
    # 3.1.2 核心需求与赛场数据闭环
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '3.1.2  核心需求与赛场数据闭环', level=1)

    add_heading_styled(doc, '数据闭环链路', level=2)
    add_para(doc, '系统完整数据流如下：')

    add_styled_table(doc,
        ['环节', '传感器/模块', '输出', '下游消费者'],
        [
            ['① 图像采集', '海康 MV-CS060 6MP @60FPS', '/image_raw (零拷贝)', 'Detect (YOLO)'],
            ['② 目标检测', 'YOLO TensorRT 推理', '/detect_result (装甲板+3D框)', 'Resolve'],
            ['③ 坐标解算', 'Resolve (PnP+HeightGrid)', '/resolve_result (世界坐标)', 'radar_serial_node'],
            ['④ 点云采集', 'Livox Mid-70 @10Hz', '/livox/lidar', 'DynamicCloud'],
            ['⑤ 动态提取', 'DynamicCloud', '/livox/lidar_cluster (候选点)', 'Cluster'],
            ['⑥ 聚类', 'Cluster (欧式聚类)', '/livox/lidar_cluster (目标簇)', 'KalmanFilter'],
            ['⑦ 多目标跟踪', 'KalmanFilter (5槽位)', '/kalman_detect (跟踪坐标)', 'radar_serial_node'],
            ['⑧ 定位', 'Localization (CUDA ICP)', '/tf (rm_frame→livox_frame)', '全链路'],
            ['⑨ 串口上报', 'radar_serial_node', '0x0305 协议 ≤5Hz', '裁判系统小地图'],
            ['⑩ 录包', 'BagRecorder', '.mcap 文件', '赛后 Foxglove 回放'],
        ],
        col_widths=[2.2, 3.5, 4.5, 5.8]
    )

    add_heading_styled(doc, '各子系统需求-实现对照', level=2)

    add_styled_table(doc,
        ['子系统', '需求', '实现方式', '评估'],
        [
            ['视觉检测', '≥30FPS 实时检测全部敌方机器人', 'YOLO + TensorRT，detect_max_process_fps=30', '✅ 达标'],
            ['多目标跟踪', '5类目标同时跟踪', 'KalmanFilter 5槽位，含出生/死亡管理', '✅ 达标'],
            ['串口上报', '≤5Hz 上报坐标，CRC 校验', '0x0305 协议，CRC8/CRC16', '✅ 达标'],
            ['LiDAR 定位', '全场点云配准→自身定位', 'CUDA ICP + PCD 地图匹配', '✅ 达标'],
            ['自动录包', '每场自动录制，赛后回放', 'BagRecorder + loop_playback', '✅ 达标'],
            ['标定', '快速外参标定', '5点 PnP，~2分钟', '⚠️ 依赖人工点选'],
            ['夜间/室内', '暗光环境正常检测', 'camera_driver_night.yaml 手动切换', '⚠️ 需手动切换'],
        ],
        col_widths=[2.5, 4, 5, 4.5]
    )

    add_heading_styled(doc, 'LiDAR-相机融合精度（核心 KPI）', level=2)

    add_para(doc, '以下是基于对手专项包，使用 analyze_lidar_camera_overlap.py 工具在 120ms 时间窗口、0.5m/0.8m/1.0m 三个阈值下的融合精度分析：')

    add_styled_table(doc,
        ['对手', '总视觉点数', 'LiDAR匹配率', '≤0.5m重叠率', '≤1.0m重叠率', '中位误差'],
        [
            ['浙大1（红）', '3,903', '88.3%', '13.5%', '20.5%', '2.69m'],
            ['浙大2（红）', '9,250', '89.7%', '29.5%', '41.5%', '1.14m'],
            ['中国石油大学（红）', '77', '53.2%', '27.3%', '28.6%', '0.32m'],
        ],
        col_widths=[3.5, 2.5, 2.5, 2.5, 2.5, 2.5]
    )

    add_rich_para(doc, [
        ('★ 关键发现：', True, RGBColor(0xC0, 0x39, 0x2B)),
        ('浙大2 的融合精度显著优于浙大1（≤1.0m 重叠率 41.5% vs 20.5%，中位误差 1.14m vs 2.69m），说明标定质量/场地条件/对手机器人尺寸对融合精度影响巨大。中国石油大学仅 77 个检测点且匹配率仅 53.2%，暴露了系统在"低检测密度"场景下的脆弱性。', False, None),
    ])

    add_heading_styled(doc, '各兵种融合精度差异（以浙大2数据为例）', level=2)

    add_styled_table(doc,
        ['兵种', '视觉点数', '≤0.5m重叠率', '≤1.0m重叠率', '中位误差', '评价'],
        [
            ['Hero', '1,423', '45.1%', '54.6%', '0.46m', '🟢 最佳，体型大、点云丰富'],
            ['Infantry 3', '1,511', '33.5%', '46.2%', '1.03m', '🟡 中等'],
            ['Infantry 4', '1,133', '29.4%', '40.2%', '1.34m', '🟡 中等偏下'],
            ['Engineer', '530', '22.1%', '36.4%', '1.25m', '🔴 最差，体型小、运动特殊'],
            ['Sentry', '696', '22.7%', '36.8%', '1.55m', '🔴 差，静止但遮挡多'],
        ],
        col_widths=[2.2, 2, 2.5, 2.5, 2.5, 4.3]
    )

    add_para(doc, 'Engineer 和 Sentry 的融合精度仅为 Hero 的 2/3，因为这两种机器人体积小、点云稀疏。如果串口上报对所有兵种的坐标精度要求相同，则对 Engineer/Sentry 的上报存在系统性偏差。')

    # ==========================================
    # 3.1.3 核心方案成败总结与经验沉淀
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '3.1.3  核心方案成败总结与经验沉淀', level=1)

    add_heading_styled(doc, '✅ 成功方案', level=2)

    add_styled_table(doc,
        ['方案', '技术细节', '效果', '可复用经验'],
        [
            ['ComposableNode\n零拷贝容器', 'camera/detector/resolve\n打包为 component_\ncontainer_mt 单进程', '消除图像 intra-process\nDDS 序列化开销', '实时视觉链路首选\nROS2 Component 架构'],
            ['双容器隔离', 'camera_detector_container\n∥ localization/cluster/\nkalman (LiDAR)', '两链路互不阻塞，\nGPU/CPU 资源隔离', '异构传感器\n应分容器部署'],
            ['TensorRT\n预编译引擎', '.onnx → .engine\n离线转换，启动加载', '免去首次推理的\nTRT build 耗时\n（可长达数分钟）', '赛前必须完成\n所有模型 TRT 编译'],
            ['/resolve_result\nfallback', 'KalmanFilter 丢失时\n用纯视觉PnP坐标补位', '极大提升\n目标连续性', '"保底链路"是\n比赛系统的刚需'],
            ['自动录包 +\n回放隔离', 'bag_out_matrix_*.yaml\n独立外参，bag.launch.py\n完整回放链路', '赛后复盘\n不影响比赛配置', '回放与实时\n必须完全隔离配置'],
            ['一键启动', 'match.launch.py 根据\nteam 自动选择红/蓝\n标定文件', '减少人为\n配置失误', '把"人工判断"\n写成代码逻辑'],
            ['Foxglove\n离线回放', 'bag.launch.py +\nfoxglove_bridge\nport:=8765', '3D 可视化排查\n检测/跟踪问题', '比赛实时不开，\n事后复盘必备'],
        ],
        col_widths=[2.5, 4, 3.5, 6]
    )

    add_heading_styled(doc, '❌ 不足/失败方案', level=2)

    add_styled_table(doc,
        ['方案', '问题描述', '根因分析', '改进方向'],
        [
            ['手动 PnP 标定', '每场搬动后需人工\n点5个点，耗时~2分钟', '无自动特征匹配\n或固定标记点', '引入 ArUco 码或\n固定标定杆，一键标定'],
            ['USB-TTL\n热熔胶固定', '5.21 疑似因接头松动\n导致连续6场空包', '杜邦线+热熔胶\n不是可靠方案', '换焊接+JST接插件，\n增加串口心跳检测'],
            ['单 LiDAR\n无冗余', 'Mid-70 视场角有限\n（圆形~70°），盲区大', '成本/重量限制', '评估 Mid-360（360°FOV）\n或双 Mid-70'],
            ['HeightGrid\n离线固定', '地形网格赛前生成\n一次，不可在线更新', '赛场地面可能不平', '增加在线地形估计\n（RANSAC 平面拟合）'],
            ['相机曝光\n手动切换', 'day/night 模式\n需人工改 yaml', '无环境光传感器', '增加自动曝光 ROI\n或光敏电阻'],
            ['无启动\n健康检查', '空包录制未阻止', '启动后未验证\n全部传感器就绪', '5s 全传感器心跳检测\n失败则阻止录包+告警'],
        ],
        col_widths=[2.5, 3.5, 3.5, 6.5]
    )

    add_heading_styled(doc, '构建日志揭示的研发节奏', level=2)

    add_styled_table(doc,
        ['日期', '构建次数', '峰值事件数', '特点'],
        [
            ['5.19', '1 次', '—', '赛前最终构建基线'],
            ['5.20', '21 次', '1,367 (17:51)', '🔴 比赛首日仍在高频修改，凌晨2:00-3:45有10次构建'],
            ['5.21', '33 次', '558 (14:52)', '🔴 边打边改，下午场间隙密集构建'],
            ['5.22', '23 次', '1,359 (14:27)', '赛后调优，下午两轮重量级构建'],
            ['5.23', '4 次', '1,333 (21:17)', '赛后集中修改融合算法（kalman_filter）'],
        ],
        col_widths=[2, 2.5, 3.5, 8]
    )

    add_rich_para(doc, [
        ('★ 关键教训：', True, RGBColor(0xC0, 0x39, 0x2B)),
        ('比赛首日（5.20）凌晨 2:00-3:45 仍有 10 次构建，首日下午-晚间又有 11 次构建。大量关键修改被推迟到了赛场上甚至比赛进行中——"赛前测试不充分，赛场上改代码"是高危模式，应引入"赛前 N 天代码冻结"机制。', False, None),
    ])

    # ==========================================
    # 3.1.4 人力资源与能力缺口复盘
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '3.1.4  人力资源与能力缺口复盘', level=1)

    add_heading_styled(doc, '团队能力矩阵', level=2)

    add_styled_table(doc,
        ['能力域', '自评', '证据', '缺口影响'],
        [
            ['C++/ROS2\n系统开发', '🟢 强', '46个C++文件，~22K行，\nComponent/Launch/零拷贝\n均正确使用', '—'],
            ['Python\n工具链', '🟢 较强', '18个Python文件，\n覆盖标定/超分/\n重叠分析/地图构建', '—'],
            ['AI 模型训练', '🔴 缺失', 'model/目录仅有ONNX/TRT\n推理文件，无训练脚本\n和数据集', '无法用赛场数据 fine-tune，\n对新外观机器人泛化差'],
            ['CUDA/GPU\n优化', '🟡 偏弱', '仅 localization 有 CUDA ICP，\ndetect 用 TRT 但无\n自定义 CUDA kernel', '检测链路 GPU 利用率\n优化空间大'],
            ['硬件/电气\n工程', '🔴 缺失', '无原理图/PCB设计文件，\n全部依赖成品模块', 'USB-TTL 可靠性、\nLiDAR 供电等问题反复出现'],
            ['自动化测试', '🔴 缺失', '无 CI/CD、单元测试、\n集成测试框架', '赛前 preflight 全靠\n人肉 checklist'],
            ['远程运维\n/监控', '🔴 缺失', '无远程日志上报、\n无健康监控面板', '赛场故障只能\n现场排查'],
        ],
        col_widths=[2, 1.5, 5, 7.5]
    )

    add_heading_styled(doc, '核心人力缺口排序（按赛场影响）', level=2)

    add_styled_table(doc,
        ['优先级', '缺口', '赛场影响', '建议'],
        [
            ['P0', 'AI 模型训练', '无法针对对手机器人外观做 domain adaptation', '下赛季招募 ML 方向队员或与 AI 实验室合作'],
            ['P0', '硬件工程', '串口/供电可靠性的根因在物理层，软件无法根治', '招募硬件/电子方向队员，至少做到接插件规范化'],
            ['P1', '自动化测试', 'preflight_check.md 只是 checklist，无自动化验证', '建立基于 rosbag 的回归测试框架'],
            ['P1', '远程监控', '无法在候场区/观众席监控雷达状态', '开发轻量 Web 仪表盘'],
        ],
        col_widths=[1.5, 2.5, 5.5, 6.5]
    )

    # ==========================================
    # 3.1.5 赛场故障与测试盲区复盘
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '3.1.5  赛场故障与测试盲区复盘', level=1)

    add_heading_styled(doc, '已知故障模式及发生情况', level=2)

    add_styled_table(doc,
        ['故障', '根因', '5.20-5.24\n发生推测', '当前检测方式', '自动化程度'],
        [
            ['相机打不开', 'USB 3.0 接触/供电', '高频（≥3次）', '自动重试 + lsusb', '🟢 半自动'],
            ['裁判串口无点位', 'USB-TTL 松动', '中频\n(5.21 6连空包)', 'ls /dev/gimbal', '🟡 仅存在性检查'],
            ['LiDAR 无数据', '网线/供电', '中频', 'ping 192.168.1.1', '🟡 仅连通性检查'],
            ['3D 框歪斜', '搬动后外参失效', '每场搬动后必现', '肉眼判断', '🔴 无自动检测'],
            ['坐标全反', 'team 红蓝设反', '低频', 'warn_if_same_side', '🟡 启动时检查'],
            ['进程残留', '异常退出', '低频', '无', '🔴 需手动 pkill'],
        ],
        col_widths=[3, 3, 2.5, 3.5, 4]
    )

    add_heading_styled(doc, '测试盲区识别', level=2)

    add_styled_table(doc,
        ['盲区', '风险等级', '触发场景', '建议检测方案'],
        [
            ['环境光突变', '🔴 高', '室内灯光切换/窗帘开合/阳光直射', '图像平均亮度监控→自动切换曝光'],
            ['多 LiDAR 互干扰', '🟡 中', '多支队伍雷达同时工作', '赛前多机联合测试（需协调）'],
            ['长时运行内存泄漏', '🟡 中', '连续 BO5（~1小时持续运行）', '内存 RSS 监控+超阈值自动重启'],
            ['串口 CRC 错误累积', '🟡 中', '电磁干扰/线缆老化', 'CRC 错误率统计+超阈值告警'],
            ['GPU 降频/过热', '🟡 中', '高温环境长时间推理', 'nvidia-smi 温度/频率监控'],
            ['磁盘写满', '🟢 低', '大量录包未清理', '磁盘剩余空间监控'],
            ['裁判系统协议变更', '🟡 中', '官方固件升级', '协议版本号校验'],
        ],
        col_widths=[3, 2, 4.5, 6.5]
    )

    add_rich_para(doc, [
        ('★ 核心问题：', True, RGBColor(0xC0, 0x39, 0x2B)),
        ('5月20日全天 21 次构建，其中凌晨 2:00-3:45 有 10 次。这些凌晨构建说明：① 赛前最后一天仍在改核心代码；② 没有"代码冻结"机制；③ 凌晨修改未经充分测试就用于当天比赛。这是整个赛季最大的流程风险。', False, None),
    ])

    # ==========================================
    # 3.1.6 遗留问题与技术传承库
    # ==========================================
    doc.add_page_break()
    add_heading_styled(doc, '3.1.6  遗留问题与技术传承库', level=1)

    add_heading_styled(doc, '遗留问题清单（按优先级排序）', level=2)

    add_styled_table(doc,
        ['编号', '问题', '严重度', '当前状态', '建议方案', '预估工时'],
        [
            ['L-01', '手动标定流程', '🔴 高', '每次搬动耗2分钟\n+人为误差', '场地固定标记点\n+ArUco自动标定', '2 周'],
            ['L-02', 'USB-TTL 可靠性', '🔴 高', '热熔胶+杜邦线\n5.21出现连续故障', '换焊接+JST接插件\n增加串口心跳检测', '3 天'],
            ['L-03', '无启动健康检查', '🔴 高', '空包照录不误', 'launch后全传感器\n心跳检测，失败则\n阻止录包+语音告警', '2 天'],
            ['L-04', 'Engineer/Sentry\n融合精度低', '🟡 中', '≤1m重叠仅36-37%', '优化小目标聚类参数\n增加点云累积帧数', '1 周'],
            ['L-05', 'AI模型无法自训练', '🟡 中', '仅有推理文件', '建立训练pipeline\n+对手数据集标注', '3 周'],
            ['L-06', '无自动化测试', '🟡 中', '仅preflight_check.md', '基于rosbag的\n回归测试框架', '2 周'],
            ['L-07', '远程监控缺失', '🟡 中', '必须现场看屏幕', '轻量Web仪表盘\n(CPU/GPU/检测数\n/串口状态)', '1 周'],
            ['L-08', '单LiDAR无冗余', '🟢 低', 'Mid-70单点故障', '评估Mid-360替代', '调研'],
            ['L-09', 'HeightGrid\n离线固定', '🟢 低', '赛场地面变化\n无法适应', '在线RANSAC\n平面拟合', '3 天'],
            ['L-10', '相机曝光手动切换', '🟢 低', '需人工改yaml', '图像亮度反馈\n自动切换', '1 天'],
        ],
        col_widths=[1.2, 3, 1.5, 3.5, 4.3, 2.5]
    )

    add_heading_styled(doc, '技术传承库（新人接手必读）', level=2)

    add_styled_table(doc,
        ['文档', '路径', '内容', '完整度'],
        [
            ['启动手册', 'README.md', '比赛启动命令+故障排查', '✅ 完整'],
            ['详细说明', 'README_DETAIL.md', '接线/标定/比赛日流程/回放', '✅ 完整'],
            ['赛前测试', 'docs/preflight_check.md', '系统链路测试checklist', '✅ 完整'],
            ['雷达配置', 'config/radar_runtime.yaml', '全部可调参数+注释', '✅ 完整'],
            ['融合精度报告', 'reports/*.txt', '各对手LiDAR-相机重叠率', '✅ 有数据'],
            ['构建历史', 'log/build_*/', '84次完整构建日志', '✅ 完整'],
            ['比赛录包', 'bags/match_*/', '59场原始数据', '✅ 完整'],
            ['架构设计文档', '—', '系统架构/API文档', '🔴 缺失'],
            ['模型训练指南', '—', '训练pipeline/数据集说明', '🔴 缺失'],
            ['硬件接线图', '—', '接线图/供电拓扑', '🔴 缺失'],
        ],
        col_widths=[3, 4.5, 5, 3.5]
    )

    add_heading_styled(doc, '对 2027 赛季的核心建议', level=2)

    suggestions = [
        ('1. 代码冻结日', '赛前至少 3 天锁定核心链路代码，仅允许修复 P0 级 bug，所有修改必须通过自动化回归测试。'),
        ('2. 自动化赛前检查', '将 preflight_check.md 变成自动化脚本，每个检查项返回 PASS/FAIL，任意 FAIL 则阻止比赛启动。'),
        ('3. 串口冗余设计', '硬件层面根治 USB-TTL 可靠性问题——换用工业级接插件、增加串口心跳包检测、引入自动重连。'),
        ('4. 建立对手模型库', '用 2026 赛季采集的 3 个对手（浙大1、浙大2、中国石油大学）数据标注并训练专属检测模型，提升对已知对手的检测精度。'),
        ('5. 补充架构文档', '编写系统架构设计文档和 API 文档，让新人能在一周内理解全链路，而非靠逆向阅读 ~22K 行代码。'),
        ('6. 引入 CI/CD', '基于 GitHub Actions 或本地 Jenkins，每次提交自动编译+运行 rosbag 回归测试，杜绝"赛前改崩了不知道"。'),
    ]
    for title, desc in suggestions:
        add_rich_para(doc, [
            (title + '：', True, RGBColor(0x1F, 0x4E, 0x79)),
            (desc, False, None),
        ])

    # ========== 保存 ==========
    output_path = '/home/zst/T/reports/T雷达站_RMUC2026赛季复盘报告.docx'
    doc.save(output_path)
    print(f'✅ 报告已生成: {output_path}')

if __name__ == '__main__':
    main()
